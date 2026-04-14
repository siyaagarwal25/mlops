from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).parent
DIAGRAM_PATH = ROOT / "mlops_architecture_diagram.png"
DOCX_PATH = ROOT / "Diabetes_MLOps_Project_Report.docx"


def create_architecture_diagram(output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(14, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis("off")

    boxes = [
        (0.6, 5.8, 2.2, 1.1, "Data\nSource"),
        (3.4, 5.8, 2.2, 1.1, "Model\nTraining\n(train.py)"),
        (6.2, 5.8, 2.2, 1.1, "Model\nArtifact\n(diabetes_model.pkl)"),
        (9.0, 5.8, 2.2, 1.1, "FastAPI\nService\n(main.py)"),
        (11.8, 5.8, 2.0, 1.1, "Monitoring\n/metrics\nPrometheus"),
        (3.4, 3.5, 2.4, 1.1, "Tests\n(pytest)") ,
        (6.4, 3.5, 2.4, 1.1, "Docker\nContainer"),
        (9.4, 3.5, 2.8, 1.1, "Kubernetes\nDeployment\n(HPA, Probes, PDB)"),
        (6.4, 1.2, 3.2, 1.1, "GitHub Actions\nCI/CD Pipeline"),
    ]

    for x, y, w, h, label in boxes:
        rect = plt.Rectangle((x, y), w, h, fill=False, linewidth=1.8)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=10)

    arrows = [
        ((2.8, 6.35), (3.4, 6.35)),
        ((5.6, 6.35), (6.2, 6.35)),
        ((8.4, 6.35), (9.0, 6.35)),
        ((11.2, 6.35), (11.8, 6.35)),
        ((7.3, 5.8), (7.3, 4.6)),
        ((10.2, 5.8), (10.8, 4.6)),
        ((4.6, 3.5), (7.0, 2.3)),
        ((7.6, 3.5), (7.8, 2.3)),
        ((10.8, 3.5), (8.8, 2.3)),
    ]

    for start, end in arrows:
        ax.annotate(
            "",
            xy=end,
            xytext=start,
            arrowprops=dict(arrowstyle="->", lw=1.5),
        )

    ax.text(0.6, 0.35, "End-to-End MLOps Workflow: Train -> Serve -> Test -> Containerize -> Deploy -> Monitor -> Automate", fontsize=10)

    fig.tight_layout()
    fig.savefig(output_path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def create_docx(diagram_path: Path, output_path: Path) -> None:
    doc = Document()

    title = doc.add_heading("Diabetes Prediction MLOps Project Report", level=0)
    title.alignment = 1

    subtitle = doc.add_paragraph(
        f"Generated on: {datetime.now().strftime('%d %B %Y')}\n"
        "Capstone Topic: DevOps for AI (Unit V)"
    )
    subtitle.alignment = 1

    add_heading(doc, "1. Project Overview", level=1)
    doc.add_paragraph(
        "This project predicts whether a patient is diabetic and shows how to deploy an AI model using "
        "DevOps practices. The goal is to make the model easy to train, test, deploy, monitor, and maintain."
    )

    add_heading(doc, "2. Open-Source Tools Used", level=1)
    for item in [
        "Scikit-learn for model training",
        "FastAPI for model serving",
        "Pytest for automated testing",
        "Docker for containerization",
        "Kubernetes (Kind) for orchestration",
        "GitHub Actions for CI/CD",
        "Prometheus metrics for monitoring",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Architecture Diagram", level=1)
    doc.add_paragraph(
        "The diagram below shows the complete MLOps workflow from data and training to deployment, monitoring, and automation."
    )
    doc.add_picture(str(diagram_path), width=Inches(6.8))
    doc.add_paragraph("Figure: End-to-end MLOps architecture")

    add_heading(doc, "4. Implementation in Simple Steps", level=1)
    steps = [
        "Train model using train.py and save it as diabetes_model.pkl.",
        "Load the model in FastAPI (main.py) and expose prediction endpoint /predict.",
        "Add health endpoints (/health and /ready) for deployment checks.",
        "Add monitoring endpoint (/metrics) with Prometheus metrics.",
        "Write test cases in test_main.py and enforce coverage with pytest.ini.",
        "Build container image using Dockerfile and run app in a non-root user.",
        "Deploy app on Kubernetes using k8s-deploy.yml with probes, HPA, and PDB.",
        "Automate checks and builds in .github/workflows/ci-cd.yml.",
    ]
    for idx, step in enumerate(steps, start=1):
        doc.add_paragraph(f"{idx}. {step}")

    add_heading(doc, "5. Key DevOps Features Implemented", level=1)
    for item in [
        "Automated tests with coverage threshold",
        "Code quality checks (Black and Flake8)",
        "Security scan using Bandit",
        "Container healthcheck and secure runtime",
        "Kubernetes readiness/liveness/startup probes",
        "Auto-scaling using Horizontal Pod Autoscaler",
        "Availability protection using PodDisruptionBudget",
        "Operational metrics for traffic, latency, errors, and model readiness",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. How to Demonstrate to Professor", level=1)
    demo_steps = [
        "Run python train.py",
        "Run uvicorn main:app --reload --host 0.0.0.0 --port 8000",
        "Show /docs, /health, /ready, /metrics endpoints",
        "Run pytest and show all tests passing with coverage",
        "Build and run Docker image",
        "Deploy to Kind/Kubernetes and show kubectl get deploy,po,svc,hpa,pdb",
        "Show successful GitHub Actions pipeline run",
    ]
    for item in demo_steps:
        add_bullet(doc, item)

    add_heading(doc, "7. Conclusion", level=1)
    doc.add_paragraph(
        "This project successfully demonstrates DevOps for AI using an open-source stack. "
        "It includes model lifecycle, API deployment, testing, CI/CD, containerization, orchestration, and monitoring."
    )

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.save(output_path)


if __name__ == "__main__":
    create_architecture_diagram(DIAGRAM_PATH)
    create_docx(DIAGRAM_PATH, DOCX_PATH)
    print(f"Created diagram: {DIAGRAM_PATH}")
    print(f"Created report: {DOCX_PATH}")
